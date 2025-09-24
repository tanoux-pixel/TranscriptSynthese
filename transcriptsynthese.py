# ==============================================================
# TRANSCRIPTEUR & SYNTHÈSE IA AVANCÉE - SANTÉ / QUALITÉ / DOC
# --------------------------------------------------------------
# Ce script tout-en-un permet :
# 1. Transcription automatique (audio, vidéo, YouTube → texte) avec Whisper.
# 2. Menu avancé pour paramétrer la synthèse IA (public, contexte, style, extraction d'éléments…).
# 3. Génération de prompt IA avancé.
# 4. Synthèse IA (Mistral local par défaut, facilement modifiable pour OpenAI/Gemini).
# 5. Export Word (.docx) avec structuration.
# 6. Double usage : terminal (menu interactif) ou API (FastAPI).
# --------------------------------------------------------------
# DEPENDANCES : torch, whisper, yt-dlp, requests, docx (python-docx),
# fastapi, uvicorn, transformers, accelerate
# ==============================================================

import os
import sys
import time
import json
import torch
from pathlib import Path

# Pour transcription audio/vidéo/YouTube
try:
    import whisper
    from yt_dlp import YoutubeDL
except ImportError:
    print("[!] Installer whisper et yt-dlp : pip install whisper yt-dlp")
    sys.exit(1)

# Pour l'export Word
try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("[!] Installer python-docx : pip install python-docx")
    sys.exit(1)

# Pour Mistral local
try:
    from transformers import AutoTokenizer, AutoModelForCausalLM
except ImportError:
    print("[!] Installer transformers : pip install transformers accelerate")
    sys.exit(1)

# Pour le mode API (optionnel)
try:
    from fastapi import FastAPI, File, UploadFile
    from fastapi.responses import FileResponse
    import uvicorn
except ImportError:
    pass  # Pas bloquant pour l'usage menu classique

# ===============================
# PARAMÈTRES GÉNÉRIQUES (menus)
# ===============================

PUBLICS_CIBLES = [
    "Grand public (adolescents/adultes)",
    "Familles",
    "Professionnels de santé",
    "Institutionnels (ARS, HAS…)",
    "Décideurs/gestionnaires",
    "Patients/résidents",
    "Autre (saisie manuelle)"
]

CONTEXTES = [
    "Rapport interne/qualité",
    "Support à une réunion",
    "Communication à destination des usagers",
    "Restitution à la direction",
    "Dossier de certification (HAS…)",
    "Présentation à une commission (CVS, CME…)",
    "Article ou publication externe",
    "Autre (saisie manuelle)"
]

NIVEAUX_LANGUE = [
    "Très grand public (aucun jargon)",
    "Pédagogique (explications, définitions…)",
    "Professionnel courant (style rapport de service)",
    "Technique (lexique professionnel/savant)",
    "Prêt à publier (style éditorial, sans fautes, structuré)",
    "Journalistique (style article de presse, narratif)",
    "Scientifique (style article scientifique)",
    "Administratif (style institutionnel)",
    "Autre (saisie manuelle)"
]

TONS = [
    "Autoritaire", "Engageant", "Neutre", "Pédagogique", "Autre"
]

MODELES_MISTRAL = [
    "mistralai/Mistral-7B-Instruct-v0.1",
    "mistralai/Mistral-7B-Instruct-v0.2",
    "mistralai/Mistral-7B-v0.1"
]

# Variables globales pour le cache du modèle
_cached_model = None
_cached_tokenizer = None
_cached_model_name = None

# ===============================
# OUTILS DE MENU INTERACTIF
# ===============================

def choix_menu(label, options):
    """
    Affiche un menu à choix unique, retourne le texte choisi.
    """
    print(f"\n{label}")
    for i, opt in enumerate(options, 1):
        print(f"{i}. {opt}")
    choix = input("Votre choix (numéro ou texte) : ").strip()
    if choix.isdigit() and 1 <= int(choix) <= len(options):
        val = options[int(choix) - 1]
        if "autre" in val.lower():
            return input("Veuillez préciser : ")
        else:
            return val
    else:
        return choix

def choix_multi(label, options):
    """
    Menu à choix multiples, retourne une liste booléenne (True pour sélectionné).
    """
    print(f"\n{label} (séparez les numéros par une virgule pour multi-sélection)")
    for i, opt in enumerate(options, 1):
        print(f"{i}. {opt}")
    choix = input("Votre choix : ").strip().replace(" ", "")
    idxs = [int(x)-1 for x in choix.split(",") if x.isdigit() and 1 <= int(x) <= len(options)]
    return [i in idxs for i in range(len(options))]

# ===============================
# TRANSCRIPTION AUDIO/VIDEO
# ===============================

def is_youtube_url(url):
    """Détecte si l'entrée est une URL YouTube."""
    import re
    youtube_regex = r'(https?://)?(www\.)?(youtube\.com|youtu\.be)/'
    return re.match(youtube_regex, url) is not None

def download_youtube_audio(url):
    """
    Télécharge l'audio d'une vidéo YouTube au format FLAC (qualité pour transcription).
    Retourne le nom de fichier local.
    """
    print(f"\n[YouTube] Téléchargement de l'audio depuis : {url}")
    ydl_opts = {
        'format': 'bestaudio/best',
        'outtmpl': 'yt_audio_%(id)s.%(ext)s',
        'noplaylist': True,
        'quiet': True,
        'postprocessors': [{
            'key': 'FFmpegExtractAudio',
            'preferredcodec': 'flac',
            'preferredquality': '0',
        }]
    }
    with YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=True)
        audio_file = f"yt_audio_{info['id']}.flac"
        print(f"[YouTube] Audio extrait : {audio_file}")
        return audio_file

def transcrire_fichier_audio(fichier, modele="base"):
    """
    Utilise Whisper pour transcrire un fichier audio/vidéo.
    """
    device = "cuda" if torch.cuda.is_available() else "cpu"
    print(f"[Whisper] Chargement du modèle {modele} sur {device}...")
    model = whisper.load_model(modele, device=device)
    print("[Whisper] Transcription en cours...")
    result = model.transcribe(fichier)
    texte = result["text"]
    nom_sortie = Path(fichier).stem + "_transcription.txt"
    with open(nom_sortie, "w", encoding="utf-8") as f:
        f.write(texte)
    print(f"[Whisper] Transcription sauvegardée : {nom_sortie}")
    return texte, nom_sortie

# ===============================
# PROMPT BUILDER AVANCÉ
# ===============================

def build_prompt(texte, params):
    """
    Génère le prompt complet pour la synthèse IA selon les paramètres utilisateur.
    """
    public = params["public"]
    contexte = params["contexte"]
    style = params["niveau_langue"]
    ton = params["ton"]
    objectifs = params["objectifs"]
    synthese_types = params["synthese_types"]  # liste de bools
    structuration = params["structuration"]    # liste de bools

    prompt = (
        f"Tu es un assistant IA expert en synthèse et vulgarisation en santé.\n"
        f"Ta mission : produire une synthèse fidèle, claire, structurée et utile du texte suivant, "
        f"en français, à destination du public suivant : **{public}**.\n"
        f"Contexte d'usage : **{contexte}**.\n"
        f"Niveau de langue : **{style}**.\n"
        f"Ton : **{ton}**.\n"
        f"Objectifs : {objectifs}.\n\n"
        f"Consignes :\n"
    )

    # Types de synthèse (résumé court, structuré, analytique, critique…)
    instructions = []
    if synthese_types[0]:
        instructions.append("- Commence par un résumé ultra-court (3 phrases max, va à l'essentiel).")
    if synthese_types[1]:
        instructions.append("- Ajoute un résumé structuré avec titres/sous-titres.")
    if synthese_types[2]:
        instructions.append("- Rédige une synthèse analytique : structure avec un plan, arguments, points forts/faibles.")
    if synthese_types[3]:
        instructions.append("- Intègre une analyse critique (mais reste neutre et factuel).")
    if synthese_types[4]:
        instructions.append("- Combine les différents styles ci-dessus si pertinent.")
    if synthese_types[5]:
        instructions.append("- Génère plusieurs versions de synthèse différentes si pertinent.")

    # Structuration spécifique
    if structuration[0]:
        instructions.append("- Structure la synthèse selon les chapitres/parties du texte d'origine si présents.")
    if structuration[1]:
        instructions.append("- Regroupe les idées par thématique.")
    if structuration[2]:
        instructions.append("- Mets en avant les citations ou éléments factuels précis entre guillemets.")

    # Extraction d'éléments spécifiques
    instructions += [
        "- En fin de synthèse, liste les points clés (bullet points).",
        "- Fournis une table des matières détaillée et numérotée.",
        "- Rédige un glossaire expliquant chaque notion importante rencontrée.",
        "- Sélectionne les citations marquantes entre guillemets.",
        "- Propose des recommandations pratico-pratiques issues de la transcription, ou des recommandations reconnues dans le monde de la santé en France.",
    ]
    # Bloc "Synthèse augmentée"
    instructions += [
        "- En fin de synthèse, ajoute un bloc d'analyse critique IA, neutre, sans prise de position.",
        "- Recherche toutes les références juridiques, lois, articles cités, et génère un résumé automatique de chaque source.",
        "- Dresse une liste complète des éléments juridiques détectés.",
        "- Ajoute un encadré 'Encadré juridique' distinct avec ces informations.",
    ]

    prompt += "\n".join(instructions)
    prompt += (
        "\n\nVoici le texte à synthétiser :\n"
        f"{texte[:12000]}"  # coupe pour éviter overflow du prompt
    )
    return prompt

# ===============================
# IA : MISTRAL LOCAL
# ===============================

def load_mistral_model(model_name="mistralai/Mistral-7B-Instruct-v0.1"):
    """
    Charge le modèle Mistral et le tokenizer. Utilise un cache pour éviter de recharger.
    """
    global _cached_model, _cached_tokenizer, _cached_model_name
    
    if _cached_model is not None and _cached_model_name == model_name:
        return _cached_tokenizer, _cached_model
    
    print(f"[Mistral] Chargement du modèle {model_name}...")
    
    try:
        tokenizer = AutoTokenizer.from_pretrained(model_name)
        
        # Configuration pour optimiser l'utilisation mémoire
        model_kwargs = {
            "torch_dtype": torch.float16,
            "device_map": "auto",
        }
        
        # Ajouter la quantification si disponible pour économiser la RAM
        try:
            model_kwargs["load_in_8bit"] = True
        except:
            pass
        
        model = AutoModelForCausalLM.from_pretrained(model_name, **model_kwargs)
        
        # Mettre en cache
        _cached_model = model
        _cached_tokenizer = tokenizer
        _cached_model_name = model_name
        
        print(f"[Mistral] Modèle chargé avec succès !")
        return tokenizer, model
        
    except Exception as e:
        print(f"[!] Erreur lors du chargement de Mistral : {e}")
        print("[!] Vérifiez que transformers et accelerate sont installés")
        sys.exit(1)

def generate_synthese_mistral(prompt, model_name="mistralai/Mistral-7B-Instruct-v0.1", temperature=0.5, max_tokens=2048):
    """
    Génère une synthèse en utilisant Mistral local via Transformers.
    """
    tokenizer, model = load_mistral_model(model_name)
    
    print("[Mistral] Génération de la synthèse en cours...")
    
    # Préparer l'input avec le format chat de Mistral
    messages = [{"role": "user", "content": prompt}]
    
    try:
        # Appliquer le template de chat
        inputs = tokenizer.apply_chat_template(
            messages, 
            return_tensors="pt", 
            add_generation_prompt=True
        )
        
        # S'assurer que les inputs sont sur le bon device
        inputs = inputs.to(model.device)
        
        # Générer la réponse
        with torch.no_grad():
            outputs = model.generate(
                inputs,
                max_new_tokens=max_tokens,
                temperature=temperature,
                do_sample=True,
                pad_token_id=tokenizer.eos_token_id,
                eos_token_id=tokenizer.eos_token_id,
                repetition_penalty=1.1
            )
        
        # Décoder la réponse complète
        full_response = tokenizer.decode(outputs[0], skip_special_tokens=True)
        
        # Extraire seulement la partie générée (après le prompt)
        prompt_encoded = tokenizer.decode(inputs[0], skip_special_tokens=True)
        if prompt_encoded in full_response:
            synthese = full_response.replace(prompt_encoded, "").strip()
        else:
            synthese = full_response
        
        print(f"[Mistral] Synthèse générée ({len(synthese)} caractères)")
        return synthese
        
    except Exception as e:
        print(f"[!] Erreur lors de la génération : {e}")
        return f"Erreur lors de la génération de la synthèse : {e}"

# ===============================
# EXPORT WORD (.docx)
# ===============================

def export_word(texte_synthese, nom_sortie="synthese_finale.docx"):
    """
    Génère un fichier Word à partir de la synthèse texte.
    """
    doc = Document()
    doc.add_heading("Synthèse IA de la transcription", 0)
    
    # Ajouter métadonnées
    doc.add_paragraph(f"Généré le : {time.strftime('%d/%m/%Y à %H:%M')}")
    doc.add_paragraph("Outil : TranscripteurSynthèse IA avec Mistral local")
    doc.add_paragraph("")  # Ligne vide
    
    # Traiter le texte de synthèse par blocs
    for bloc in texte_synthese.split("\n\n"):
        if bloc.strip():
            p = doc.add_paragraph(bloc.strip())
            p.style.font.size = Pt(11)
    
    doc.save(nom_sortie)
    print(f"✅ Fichier Word généré : {nom_sortie}")

# ===============================
# WORKFLOW PRINCIPAL (TERMINAL)
# ===============================

def main():
    print("\n=== TRANSCRIPTEUR & SYNTHÈSE IA AVANCÉE (Mistral Local) ===")

    # --- 1. Sélection de la source (audio/vidéo/YouTube ou .txt déjà prêt) ---
    print("\n[1] Choisissez le fichier à traiter :")
    chemin = input("Chemin du fichier audio/vidéo/texte OU URL YouTube : ").strip()
    if is_youtube_url(chemin):
        fichier_audio = download_youtube_audio(chemin)
        source_type = "audio"
    elif Path(chemin).suffix.lower() in [".mp3", ".wav", ".flac", ".aac", ".ogg", ".m4a", ".mp4", ".mkv", ".avi", ".mov"]:
        fichier_audio = chemin
        source_type = "audio"
    elif Path(chemin).suffix.lower() in [".txt"]:
        fichier_texte = chemin
        source_type = "texte"
    else:
        print("Format non reconnu. Abandon.")
        sys.exit(1)

    # --- 2. Transcription si besoin ---
    if source_type == "audio":
        print("\n[2] Choix du modèle Whisper (1: tiny, 2: base, 3: small, 4: medium, 5: large)")
        choix = input("Votre choix (numéro, défaut=2/base) : ").strip() or "2"
        modeles_whisper = ["tiny", "base", "small", "medium", "large"]
        modele = modeles_whisper[int(choix)-1]
        texte, fichier_texte = transcrire_fichier_audio(fichier_audio, modele=modele)
    else:
        with open(fichier_texte, "r", encoding="utf-8") as f:
            texte = f.read()

    # --- 3. Menu de paramétrage synthèse IA avancée ---
    print("\n[3] Paramétrage de la synthèse IA :")
    params = {}
    params["public"] = choix_menu("Public cible :", PUBLICS_CIBLES)
    params["contexte"] = choix_menu("Contexte d'usage :", CONTEXTES)
    params["niveau_langue"] = choix_menu("Niveau de langue :", NIVEAUX_LANGUE)
    params["ton"] = choix_menu("Ton à adopter :", TONS)
    params["objectifs"] = "Informer de façon neutre, faire un rappel de la législation si possible, lister des actions concrètes si possible (issues des cas cités dans la transcription)."
    params["synthese_types"] = choix_multi(
        "Types de synthèse (1: Résumé court, 2: Résumé structuré, 3: Synthèse analytique, 4: Analyse critique, 5: Combinaison de styles, 6: Plusieurs synthèses différentes)",
        ["Résumé court", "Résumé structuré", "Synthèse analytique", "Analyse critique", "Combinaison de styles", "Plusieurs synthèses différentes"])
    params["structuration"] = choix_multi(
        "Structuration (1: Chapitres/parties du texte, 2: Par thématique, 3: Citations/faits précis)",
        ["Chapitre/parties", "Thématique", "Citations/faits"])

    # --- 4. Paramètres IA ---
    print("\n[4] Configuration Mistral :")
    print("Modèles disponibles :")
    for i, modele in enumerate(MODELES_MISTRAL, 1):
        print(f"{i}. {modele}")
    choix_modele = input("Choix du modèle (numéro, défaut=1) : ").strip() or "1"
    modele_ia = MODELES_MISTRAL[int(choix_modele)-1]
    
    try:
        temperature = float(input("Température (0=factuel, 1=créatif) [0.5] : ") or "0.5")
    except:
        temperature = 0.5
    try:
        max_tokens = int(input("Nombre max de tokens [2048] : ") or "2048")
    except:
        max_tokens = 2048

    # --- 5. Prompt builder ---
    prompt = build_prompt(texte, params)
    print("\n--- Aperçu du prompt généré ---\n")
    print(prompt[:1000], "...\n")

    # --- 6. Génération de la synthèse IA ---
    print("[IA] Génération de la synthèse avec Mistral local...")
    synthese = generate_synthese_mistral(prompt, model_name=modele_ia, temperature=temperature, max_tokens=max_tokens)
    print("[IA] Synthèse générée (début) :\n", synthese[:1000], "...\n")

    # --- 7. Export Word ---
    export_word(synthese)

    print("\n=== FIN ===\n")

# ===============================
# API FASTAPI (MODE OPTIONNEL)
# ===============================

def run_api():
    """
    Lance l'API FastAPI permettant de piloter la synthèse à distance.
    """
    app = FastAPI(title="TranscripteurSynthèse IA API", version="2.1")

    @app.post("/synthese/")
    async def synthese_api(
        file: UploadFile = File(...),
        public: str = "Grand public",
        contexte: str = "Rapport interne/qualité",
        niveau_langue: str = "Très grand public",
        ton: str = "Neutre",
        synthese_types: str = "1,2,3",
        structuration: str = "1,2",
        modele: str = "mistralai/Mistral-7B-Instruct-v0.1",
        temperature: float = 0.5,
        max_tokens: int = 2048
    ):
        """
        Endpoint pour générer une synthèse IA à partir d'un fichier texte.
        """
        texte = (await file.read()).decode("utf-8")
        params = {
            "public": public,
            "contexte": contexte,
            "niveau_langue": niveau_langue,
            "ton": ton,
            "objectifs": "Informer de façon neutre, faire un rappel de la législation si possible, lister des actions concrètes si possible (issues des cas cités dans la transcription).",
            "synthese_types": [str(i+1) in synthese_types.split(",") for i in range(6)],
            "structuration": [str(i+1) in structuration.split(",") for i in range(3)]
        }
        prompt = build_prompt(texte, params)
        synthese = generate_synthese_mistral(prompt, model_name=modele, temperature=temperature, max_tokens=max_tokens)
        nom_sortie = "synthese_api.docx"
        export_word(synthese, nom_sortie)
        return FileResponse(nom_sortie, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=nom_sortie)

    @app.get("/")
    def root():
        return {"message": "TranscripteurSynthèse IA API - Version Mistral", "endpoints": ["/synthese/"]}

    print("Lancement de l'API (FastAPI, /synthese/)...")
    uvicorn.run(app, host="0.0.0.0", port=8000)

# ===============================
# POINT D'ENTRÉE
# ===============================

if __name__ == "__main__":
    if "--api" in sys.argv:
        run_api()
    else:
        main()
